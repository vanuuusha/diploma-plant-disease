#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Конвертер diplom.md -> diplom.docx.

Оформление по ГОСТ 7.32-2017 и styles.md (Список A):
    - шрифт Times New Roman 14pt, кегль полуторный (A-2, A-3);
    - поля: 20/20/30/15 мм top/bottom/left/right (A-1);
    - абзацный отступ 1.25 см, выравнивание по ширине (A-4, A-5);
    - структурные разделы (РЕФЕРАТ, ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ, ПЕРЕЧЕНЬ
      СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ, СОДЕРЖАНИЕ, ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ,
      СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ, ПРИЛОЖЕНИЕ А...) — прописные,
      по центру, НЕ жирные, с новой страницы (A-7, A-9, A-9a, A-18);
    - нумерованные главы (1, 2, ...) и подразделы (1.1, 1.1.1) — с
      красной строки, по левому краю, НЕ прописные, НЕ жирные;
      главы начинают новую страницу (A-8, A-9, A-9a, A-18);
    - жирное начертание НЕ используется ни в заголовках, ни в теле
      текста, ни в таблицах — иерархия задаётся outlineLvl и размером
      шрифта, не начертанием (A-18);
    - подписи рисунков — под рисунком, по центру, без точки в конце
      (A-12, ГОСТ 7.32-2017 п. 6.5.6);
    - подписи таблиц — над таблицей, слева (A-13);
    - таблицы — TNR 12, границы, заголовок без жирного (A-18);
    - формулы — нативные OMML через pandoc, блочные по центру,
      инлайн в тексте (A-14, A-15);
    - TOC не вставляется: в заголовках выставлен outlineLvl, ручная
      сборка оглавления выполняется в Word (A-19);
    - имя выходного файла — только diplom.docx (A-20).
"""

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt, RGBColor, Emu


# -----------------------------------------------------------------------------
# Конфигурация
# -----------------------------------------------------------------------------

ROOT = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(ROOT, "диплом", "diplom.md")
FIG_DIR = os.path.join(ROOT, "диплом")  # пути в md относительно этой папки
OUT_PATH = os.path.join(ROOT, "diplom.docx")  # A-20: всегда diplom.docx, без суффиксов

# Набор структурных разделов по ГОСТ 7.32-2017 (A-7, A-9a). Сравнение
# идёт case-insensitive по тексту заголовка в md (после strip).
STRUCTURAL_SECTIONS = {
    "реферат",
    "термины и определения",
    "перечень сокращений и обозначений",
    "содержание",
    "введение",
    "заключение",
    "список использованных источников",
}


def _is_structural_section(text: str) -> bool:
    """True, если заголовок — структурный раздел ГОСТ (реферат, введение,
    заключение, список источников, приложение А/Б/...). Такие разделы
    выводятся прописными буквами по центру, нумерации не имеют.
    """
    t = text.strip().lower()
    if t in STRUCTURAL_SECTIONS:
        return True
    # Приложения: «ПРИЛОЖЕНИЕ А», «Приложение Б. Название» и т. п.
    if t.startswith("приложение "):
        return True
    return False


def _is_numbered_heading(text: str) -> bool:
    """True, если заголовок начинается с арабской цифры и пробела/точки
    (примеры: «1 Название», «1.1 Название», «1.1.1 Название»). Такие
    заголовки выводятся по левому краю с красной строки, строчными,
    без жирного (A-8, A-9a).
    """
    return bool(re.match(r"^\d+(\.\d+)*\.?\s+\S", text.strip()))

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
MATH_FONT = "Cambria Math"

BODY_SIZE = Pt(14)
CODE_SIZE = Pt(11)
TABLE_SIZE = Pt(12)

# Плейсхолдер-маркер и их перечень
PLACEHOLDER_MARKER = "not_yet"
placeholders = []  # (figure_number, caption, file_line)

# Источники из «Список использованных источников». Заполняется в convert().
# Ключ — номер источника (int), значение — полная библиографическая запись по ГОСТ.
SOURCES: "dict[int, str]" = {}

# Множество номеров источников, для которых реально встретился маркер **[N]**
# в теле текста. По нему определяем, какие сноски действительно надо создать.
USED_FOOTNOTE_IDS: "set[int]" = set()


# -----------------------------------------------------------------------------
# Замена простейших LaTeX на unicode
# -----------------------------------------------------------------------------

LATEX_REPLACEMENTS = [
    (r"\\cdot", "·"),
    (r"\\times", "×"),
    (r"\\approx", "≈"),
    (r"\\leq", "≤"),
    (r"\\geq", "≥"),
    (r"\\neq", "≠"),
    (r"\\pm", "±"),
    (r"\\to", "→"),
    (r"\\rightarrow", "→"),
    (r"\\leftarrow", "←"),
    (r"\\Rightarrow", "⇒"),
    (r"\\Leftarrow", "⇐"),
    (r"\\in", "∈"),
    (r"\\notin", "∉"),
    (r"\\subset", "⊂"),
    (r"\\cup", "∪"),
    (r"\\cap", "∩"),
    (r"\\forall", "∀"),
    (r"\\exists", "∃"),
    (r"\\infty", "∞"),
    (r"\\partial", "∂"),
    (r"\\nabla", "∇"),
    (r"\\sum", "Σ"),
    (r"\\prod", "∏"),
    (r"\\int", "∫"),
    (r"\\sqrt", "√"),
    (r"\\ln", "ln"),
    (r"\\log", "log"),
    (r"\\sin", "sin"),
    (r"\\cos", "cos"),
    (r"\\tan", "tan"),
    (r"\\exp", "exp"),
    (r"\\min", "min"),
    (r"\\max", "max"),
    (r"\\softmax", "softmax"),
    (r"\\text\{([^}]*)\}", r"\1"),
    (r"\\mathrm\{([^}]*)\}", r"\1"),
    (r"\\mathbf\{([^}]*)\}", r"\1"),
    (r"\\mathcal\{([^}]*)\}", r"\1"),
    (r"\\mathbb\{([^}]*)\}", r"\1"),
    (r"\\operatorname\{([^}]*)\}", r"\1"),
    (r"\\boldsymbol\{([^}]*)\}", r"\1"),
    (r"\\odot", "⊙"),
    (r"\\oplus", "⊕"),
    (r"\\otimes", "⊗"),
    (r"\\ldots", "…"),
    (r"\\cdots", "⋯"),
    (r"\\alpha", "α"),
    (r"\\beta", "β"),
    (r"\\gamma", "γ"),
    (r"\\delta", "δ"),
    (r"\\epsilon", "ε"),
    (r"\\varepsilon", "ε"),
    (r"\\zeta", "ζ"),
    (r"\\eta", "η"),
    (r"\\theta", "θ"),
    (r"\\iota", "ι"),
    (r"\\kappa", "κ"),
    (r"\\lambda", "λ"),
    (r"\\mu", "μ"),
    (r"\\nu", "ν"),
    (r"\\xi", "ξ"),
    (r"\\pi", "π"),
    (r"\\rho", "ρ"),
    (r"\\sigma", "σ"),
    (r"\\tau", "τ"),
    (r"\\upsilon", "υ"),
    (r"\\phi", "φ"),
    (r"\\varphi", "φ"),
    (r"\\chi", "χ"),
    (r"\\psi", "ψ"),
    (r"\\omega", "ω"),
    (r"\\Gamma", "Γ"),
    (r"\\Delta", "Δ"),
    (r"\\Theta", "Θ"),
    (r"\\Lambda", "Λ"),
    (r"\\Xi", "Ξ"),
    (r"\\Pi", "Π"),
    (r"\\Sigma", "Σ"),
    (r"\\Phi", "Φ"),
    (r"\\Psi", "Ψ"),
    (r"\\Omega", "Ω"),
    (r"\\,", " "),
    (r"\\;", " "),
    (r"\\:", " "),
    (r"\\!", ""),
    (r"\\quad", "    "),
    (r"\\qquad", "        "),
    (r"\\left", ""),
    (r"\\right", ""),
    (r"\\big", ""),
    (r"\\Big", ""),
    (r"\\bigg", ""),
    (r"\\Bigg", ""),
    (r"\\\\", "\n"),
]


def latex_to_unicode(s: str) -> str:
    """Конвертирует простые LaTeX-конструкции в unicode-текст."""
    # \frac{a}{b} -> a/b (простая замена)
    for _ in range(3):
        s = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", r"(\1)/(\2)", s)
    # \sqrt{a} -> √a
    s = re.sub(r"\\sqrt\{([^{}]*)\}", r"√(\1)", s)
    # оставшиеся команды
    for pat, rep in LATEX_REPLACEMENTS:
        s = re.sub(pat, rep, s)
    # Надстрочные/подстрочные — оставим как есть (будем обрабатывать через xml)
    return s


# -----------------------------------------------------------------------------
# Документ, стили, оформление
# -----------------------------------------------------------------------------

def setup_document():
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)

    # Дефолтный стиль
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    # Для поддержки кириллицы
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"), BODY_FONT)
    rFonts.set(qn("w:eastAsia"), BODY_FONT)

    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Включение автоматического переноса слов
    settings = doc.settings.element
    auto_hyphenation = OxmlElement("w:autoHyphenation")
    auto_hyphenation.set(qn("w:val"), "true")
    settings.append(auto_hyphenation)

    # Попросить Word обновить все поля при открытии — важно для поля TOC:
    # Word предложит пересобрать оглавление, и пользователю не нужно вручную
    # нажимать F9.
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    return doc


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_filler_page(doc):
    """Добавляет «зрительно пустую» страницу — абзац с разрывом страницы.

    Используется для мест, где автор позже вставит в Word титульный лист,
    задание или автоматическое оглавление. Страница содержит только
    невидимый параграф-маркер; разрыв страницы внутри параграфа гарантирует,
    что последующее содержимое окажется на следующей странице.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_section_page_num_start(section, start):
    """Устанавливает в sectPr раздела начальный номер страницы (w:pgNumType w:start)."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))


def _set_section_footer_page_number(section, *, size=Pt(14)):
    """Настраивает футер раздела: по центру — PAGE-поле с номером страницы.

    Шрифт — TNR указанного размера (по умолчанию 14pt по явному запросу
    автора: «размер нижнего колонтитула 14 — для сносок»). Футер отвязывается
    от предыдущего раздела, чтобы нумерация не наследовалась и не
    отображалась на страницах, относящихся к более ранним разделам
    (титульный лист, задание, реферат).
    """
    footer = section.footer
    footer.is_linked_to_previous = False

    # Используем существующий параграф (python-docx создаёт его автоматически)
    # или добавляем новый.
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    # Очищаем содержимое
    for r in list(fp.runs):
        r._element.getparent().remove(r._element)

    pf = fp.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    run = fp.add_run()
    run_set_font(run, BODY_FONT, size=size)
    r_el = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_el.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE \\* MERGEFORMAT "
    r_el.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_el.append(fld_sep)

    t = OxmlElement("w:t")
    t.text = "4"
    r_el.append(t)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_el.append(fld_end)


def start_numbered_section(doc, *, start_page=4, page_number_size=Pt(14)):
    """Создаёт новый раздел с разрывом страницы и нумерацией страниц.

    Назначение: после реферата (страница 3) начинается новый раздел с
    пустой страницей под автооглавление (страница 4) и далее — основной
    текст. В этом разделе в футере появляется номер страницы, начиная с 4;
    предыдущий раздел (титульный лист, задание, реферат) остаётся без
    номеров страниц.

    Возвращает объект Section для нового раздела.
    """
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    # Поля — те же, что и у первого раздела (A-1 в styles.md)
    new_section.top_margin = Mm(20)
    new_section.bottom_margin = Mm(20)
    new_section.left_margin = Mm(30)
    new_section.right_margin = Mm(15)

    # Отвязываем header/footer от предыдущего раздела
    new_section.header.is_linked_to_previous = False
    _set_section_footer_page_number(new_section, size=page_number_size)

    # Стартовый номер страницы в новом разделе
    _set_section_page_num_start(new_section, start_page)

    return new_section


def set_paragraph_style(p, *, alignment=None, first_line_indent=None,
                       bold=False, size=None, space_before=None, space_after=None,
                       line_spacing=None, keep_with_next=False):
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if keep_with_next:
        pf.keep_with_next = True


def run_set_font(run, name=BODY_FONT, size=None, bold=None, italic=None,
                 color=None, monospace=False, highlight=None):
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if highlight is not None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), highlight)
        rpr.append(shd)


# -----------------------------------------------------------------------------
# Парсер inline-разметки (bold, italic, code, math)
# -----------------------------------------------------------------------------

INLINE_PATTERN = re.compile(
    r"(\*\*\[\d+\](?:\[\d+\])*\*\*|"  # **[N]** или **[N][M]** — маркер(ы) сноски (legacy)
    r"\[\d+\](?:\[\d+\])*|"           # [N] или [N][M]... — маркер(ы) сноски (plain)
    r"\*\*[^*]+\*\*|"                 # **bold**
    r"\*[^*]+\*|"                     # *italic*
    r"`[^`]+`|"                       # `code`
    r"\$[^$]+\$)"                     # $math$
)

# Маркер сноски: **[N]** или **[N][M]...** (подряд несколько номеров) —
# старый формат разметки, оставлен для совместимости.
RE_FOOTNOTE_MARKER = re.compile(r"^\*\*((?:\[\d+\])+)\*\*$")

# Маркер сноски: [N] или [N][M]... — плоский формат. Срабатывает только
# если все номера внутри соответствуют записям «Список использованных
# источников» (проверяется через SOURCES). Это позволяет не спутать
# библиографическую ссылку с, например, `[TODO]`, `[черновик]` и прочими
# квадратными скобками в тексте.
RE_FOOTNOTE_MARKER_PLAIN = re.compile(r"^((?:\[\d+\])+)$")


def parse_inline(text):
    """Возвращает список кортежей (kind, content).

    Вид ``("footnote", [n1, n2, ...])`` обозначает маркер сноски на
    источник(и) по номеру(ам) из списка литературы.
    """
    pieces = []
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            pieces.append(("text", text[pos:m.start()]))
        tok = m.group(0)
        fn_match = RE_FOOTNOTE_MARKER.match(tok)
        fn_match_plain = RE_FOOTNOTE_MARKER_PLAIN.match(tok)
        if fn_match:
            # Legacy-формат **[N][M]...**.
            nums = [int(x) for x in re.findall(r"\[(\d+)\]", fn_match.group(1))]
            pieces.append(("footnote", nums))
        elif fn_match_plain:
            # Плоский формат [N][M]... Считаем маркером сноски только если
            # ВСЕ номера внутри — валидные записи списка источников. Иначе
            # оставляем как обычный текст, чтобы не превращать случайные
            # `[17]` в сноску на несуществующий источник.
            nums = [int(x) for x in re.findall(r"\[(\d+)\]", fn_match_plain.group(1))]
            if SOURCES and all(n in SOURCES for n in nums):
                pieces.append(("footnote", nums))
            else:
                pieces.append(("text", tok))
        elif tok.startswith("**") and tok.endswith("**"):
            pieces.append(("bold", tok[2:-2]))
        elif tok.startswith("*") and tok.endswith("*"):
            pieces.append(("italic", tok[1:-1]))
        elif tok.startswith("`") and tok.endswith("`"):
            pieces.append(("code", tok[1:-1]))
        elif tok.startswith("$") and tok.endswith("$"):
            pieces.append(("math", tok[1:-1]))
        else:
            pieces.append(("text", tok))
        pos = m.end()
    if pos < len(text):
        pieces.append(("text", text[pos:]))
    return pieces


def add_inline_runs(paragraph, text, *, base_size=BODY_SIZE, base_bold=False,
                    base_italic=False):
    """Добавляет runs с разметкой в абзац."""
    pieces = parse_inline(text)
    # Схлопываем пробел между словом и маркером сноски: в md-исходнике
    # «ООН **[1]**» имеет пробел, а в академической типографике верхний
    # индекс должен прилегать к слову — «ООН¹».
    cleaned = []
    for idx, (kind, content) in enumerate(pieces):
        if (kind == "text" and idx + 1 < len(pieces)
                and pieces[idx + 1][0] == "footnote"):
            content = content.rstrip(" \t")
        cleaned.append((kind, content))
    pieces = cleaned
    for kind, content in pieces:
        if kind == "text":
            r = paragraph.add_run(content)
            run_set_font(r, BODY_FONT, size=base_size, bold=base_bold,
                         italic=base_italic)
        elif kind == "bold":
            # A-18 / B-19: жирное начертание в теле текста не используется.
            # Маркер `**...**` в .md считаем наследственным оформлением и
            # выводим как обычный текст (без bold). Маркеры сносок
            # `**[N]**` обрабатываются отдельной веткой kind="footnote"
            # ещё до того, как попасть сюда.
            r = paragraph.add_run(content)
            run_set_font(r, BODY_FONT, size=base_size, bold=base_bold,
                         italic=base_italic)
        elif kind == "italic":
            r = paragraph.add_run(content)
            run_set_font(r, BODY_FONT, size=base_size, bold=base_bold,
                         italic=True)
        elif kind == "code":
            r = paragraph.add_run(content)
            run_set_font(r, CODE_FONT, size=Pt(12), bold=False, italic=False,
                         highlight="F2F2F2")
        elif kind == "math":
            add_inline_formula(paragraph, content, base_size=base_size)
        elif kind == "footnote":
            # content — список номеров источников; вставляем по одной сноске
            # на каждый. Между подряд идущими сносками Word сам добавит
            # верхние индексы через стиль FootnoteReference.
            for n in content:
                add_footnote_reference(paragraph, n, base_size=base_size)


# -----------------------------------------------------------------------------
# Блочные элементы
# -----------------------------------------------------------------------------

def add_heading(doc, level, text, *, is_chapter=False):
    """Вставляет заголовок с учётом правил styles.md (A-7, A-8, A-9, A-9a, A-18).

    Три случая:
        1. Структурный раздел (РЕФЕРАТ, ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и т. п.) —
           прописные буквы, по центру, с новой страницы, НЕ жирный.
           Не присваивается outlineLvl для РЕФЕРАТА (он не должен
           попадать в оглавление, см. A-7 исключение), прочим структурным
           разделам outlineLvl=0 выставляется.
        2. Нумерованный заголовок уровня 1 («1 Название», «2 Название»...) —
           по левому краю с красной строки, строчные, НЕ жирный, с новой
           страницы (начало главы).
        3. Нумерованный заголовок уровня 2+ («1.1 ...», «1.1.1 ...») —
           по левому краю с красной строки, строчные, НЕ жирный, без
           разрыва страницы.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format

    structural = _is_structural_section(text)
    # По умолчанию берём уровень заголовка из markdown (# → 1, ## → 2, ...).
    # Для структурных разделов, которые в .md размечены как `##` (например,
    # «## Термины и определения», «## Перечень сокращений…»), всё равно
    # присваиваем outlineLvl уровня 1, так как это полноценные структурные
    # разделы по ГОСТ.
    effective_level = 1 if structural else level
    is_top_level = (effective_level == 1)

    if structural:
        # A-7 / A-9a: структурный раздел — прописные по центру, новая страница.
        add_page_break_before(p)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(12)
        pf.line_spacing = 1.5
        pf.keep_with_next = True
        r = p.add_run(text.upper())
        run_set_font(r, BODY_FONT, size=Pt(14), bold=False)  # A-18: без жирного
    else:
        # Нумерованный заголовок основной части.
        if is_top_level or is_chapter:
            # Глава — с новой страницы.
            add_page_break_before(p)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT  # A-9a: по левому краю
        pf.first_line_indent = Cm(1.25)          # A-5: с красной строки
        pf.space_before = Pt(12) if not is_top_level else Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5
        pf.keep_with_next = True
        r = p.add_run(text)
        run_set_font(r, BODY_FONT, size=Pt(14), bold=False)  # A-18: без жирного

    # A-19: outlineLvl — иерархия уровней для ручной сборки оглавления в Word.
    # Исключение (A-7 подп.): РЕФЕРАТ не должен попадать в оглавление.
    if structural and text.strip().lower() == "реферат":
        # Для РЕФЕРАТА outlineLvl НЕ выставляем.
        pass
    else:
        pPr = p._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(max(0, min(effective_level - 1, 8))))
        pPr.append(outline)


def add_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbb = OxmlElement("w:pageBreakBefore")
    pPr.append(pbb)


def add_toc(doc):
    """Вставляет в начало документа раздел «СОДЕРЖАНИЕ» с автоматическим полем
    TOC, которое Word соберёт по всем заголовкам (outlineLvl 0..2 → уровни 1-3).

    Поле приходит с флагами:
        \\o "1-3"  — подхватить заголовки 1–3 уровня
        \\h        — сделать ссылки гиперссылками
        \\z        — скрыть табуляцию-заполнитель при печати в web-вид
        \\u        — брать уровень из outlineLvl (а не только из стилей)

    При открытии в Word будет предложено «Обновить поля» — включено через
    w:updateFields=true в settings.
    """
    # Заголовок «СОДЕРЖАНИЕ» — БЕЗ outlineLvl, чтобы сам не попал в оглавление
    p_title = doc.add_paragraph()
    pf = p_title.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.line_spacing = 1.5
    pf.keep_with_next = True
    r = p_title.add_run("СОДЕРЖАНИЕ")
    run_set_font(r, BODY_FONT, size=Pt(14), bold=False)  # A-18: без жирного

    # Абзац с полем TOC (fldChar begin → instrText → separate → placeholder → end)
    p_toc = doc.add_paragraph()
    pf = p_toc.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    run = p_toc.add_run()
    r_el = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_el.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_el.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_el.append(fld_sep)

    # Placeholder — будет виден, пока Word не обновит поле
    t = OxmlElement("w:t")
    t.text = (
        "Оглавление сформируется автоматически при открытии документа. "
        "Если этого не произошло — встаньте курсором в эту строку и нажмите F9, "
        "или выберите на ленте «Ссылки → Обновить оглавление»."
    )
    r_el.append(t)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_el.append(fld_end)

    # Принудительный переход на следующую страницу после оглавления —
    # через pageBreakBefore на первом следующем заголовке 1 уровня
    # (у всех глав уже есть page break перед ними), но чтобы «Введение»
    # тоже начиналось с новой страницы, полагаемся на add_heading.


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    add_inline_runs(p, text)


def add_list_item(doc, text, *, bullet=False, number=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    prefix = ("— " if bullet else f"{number}. ")
    r = p.add_run(prefix)
    run_set_font(r, BODY_FONT, size=BODY_SIZE)
    add_inline_runs(p, text)


def add_code_block(doc, lines, lang=""):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(1.0)
        pf.line_spacing = 1.15
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # Фон абзаца
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)
        r = p.add_run(line if line else " ")
        run_set_font(r, CODE_FONT, size=CODE_SIZE)


_FORMULA_CACHE = {}
_FORMULA_DPI = 300  # внутреннее разрешение рендера; не влияет на размер в документе


_OMML_CACHE = {}  # (latex_stripped, display_bool) -> omml XML string
_OMML_NS = {
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def _latex_to_omml(latex, *, display=False):
    """Конвертирует LaTeX-формулу в OMML-разметку через pandoc.

    Возвращает XML-строку:
    - display=False: <m:oMath>...</m:oMath>   (инлайн)
    - display=True:  <m:oMathPara>...<m:oMath>...</m:oMath></m:oMathPara>
                     (блочная, уже с оболочкой абзаца)
    При сбое возвращает None.
    """
    import subprocess
    import tempfile
    import zipfile
    import os as _os

    key = (latex.strip(), bool(display))
    if key in _OMML_CACHE:
        return _OMML_CACHE[key]

    if display:
        md = f"$$\n{latex.strip()}\n$$\n"
    else:
        md = f"Pad ${latex.strip()}$ pad\n"

    tmp_md = tempfile.NamedTemporaryFile(
        suffix=".md", delete=False, mode="w", encoding="utf-8"
    )
    tmp_md.write(md)
    tmp_md.close()
    tmp_docx = tmp_md.name + ".docx"
    omml = None
    try:
        subprocess.run(
            ["pandoc", "-f", "markdown", "-t", "docx", tmp_md.name, "-o", tmp_docx],
            check=True, capture_output=True, timeout=30,
        )
        with zipfile.ZipFile(tmp_docx) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        if display:
            m = re.search(r"<m:oMathPara\b[^>]*>.*?</m:oMathPara>", xml, flags=re.DOTALL)
        else:
            m = re.search(r"<m:oMath\b[^>]*>.*?</m:oMath>", xml, flags=re.DOTALL)
        if m:
            omml = m.group(0)
    except Exception:
        omml = None
    finally:
        for p in (tmp_md.name, tmp_docx):
            try:
                _os.unlink(p)
            except OSError:
                pass

    _OMML_CACHE[key] = omml
    return omml


def _omml_element(omml_xml):
    """Парсит OMML-строку и возвращает lxml-элемент, пригодный для вставки в
    параграф python-docx. Возвращает None при ошибке.
    """
    try:
        from lxml import etree
    except Exception:
        return None
    wrapped = (
        '<root xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"{omml_xml}</root>"
    )
    try:
        root = etree.fromstring(wrapped.encode("utf-8"))
    except Exception:
        return None
    if len(root) == 0:
        return None
    return root[0]


# -----------------------------------------------------------------------------
# Сноски (footnotes) по ГОСТ
# -----------------------------------------------------------------------------
#
# Задача: по маркеру ``**[N]**`` в md-исходнике вставить в docx нативную
# подстраничную сноску Word (postscript-numeric), в теле которой — полная
# библиографическая запись источника №N по ГОСТ 7.0.5 / 7.32-2017 (ровно так,
# как она приведена в разделе «Список использованных источников»).
#
# python-docx нативно не поддерживает сноски, поэтому footnotes-part
# (word/footnotes.xml + Relationships + Content Types) собирается вручную
# через низкоуровневое API pkg.OpcPackage.
#
# Правила оформления подстраничных сносок по ГОСТ (что именно складываем):
#   * пагинация сносок — сквозная по всему документу (w:numFmt="decimal");
#   * маркер в тексте — верхний индекс (arab цифра), реализуется стилем
#     «FootnoteReference»;
#   * маркер в теле сноски — также верхний индекс, реализуется стилем
#     «FootnoteText» + «FootnoteReference» на первой цифре;
#   * содержимое сноски — полная библ. запись один-в-один с «Список
#     использованных источников» (автор, заглавие, выходные данные, URL,
#     дата обращения).
# -----------------------------------------------------------------------------

FOOTNOTES_XML_TEMPLATE = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
</w:footnotes>
"""


def _ensure_footnote_styles(doc):
    """Добавляет в styles.xml стили FootnoteReference / FootnoteText,
    если их ещё нет. Без них Word не будет оформлять сноски верхним
    индексом / уменьшенным шрифтом.
    """
    from docx.oxml.ns import nsmap  # noqa: F401  — чтобы подтянуть алиасы
    styles_el = doc.styles.element

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def has_style(style_id):
        for s in styles_el.findall(qn("w:style")):
            if s.get(qn("w:styleId")) == style_id:
                return True
        return False

    if not has_style("FootnoteReference"):
        s = OxmlElement("w:style")
        s.set(qn("w:type"), "character")
        s.set(qn("w:styleId"), "FootnoteReference")
        name = OxmlElement("w:name")
        name.set(qn("w:val"), "footnote reference")
        s.append(name)
        rPr = OxmlElement("w:rPr")
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "superscript")
        rPr.append(vert)
        s.append(rPr)
        styles_el.append(s)

    if not has_style("FootnoteText"):
        s = OxmlElement("w:style")
        s.set(qn("w:type"), "paragraph")
        s.set(qn("w:styleId"), "FootnoteText")
        name = OxmlElement("w:name")
        name.set(qn("w:val"), "footnote text")
        s.append(name)
        pPr = OxmlElement("w:pPr")
        s.append(pPr)
        rPr = OxmlElement("w:rPr")
        # По явному запросу автора сноски оформляются тем же кеглем, что и
        # основной текст (14pt). В half-points единицах w:sz/w:szCs = 28.
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "28")
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "28")
        rPr.append(szCs)
        s.append(rPr)
        styles_el.append(s)


def _ensure_footnotes_part(doc):
    """Создаёт и регистрирует word/footnotes.xml, если его ещё нет. Возвращает
    корневой элемент ``<w:footnotes>``, в который можно добавлять ``<w:footnote>``.
    """
    from docx.opc.constants import CONTENT_TYPE as CT
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    from lxml import etree

    main_part = doc.part
    # Проверяем, не привязан ли уже footnotes-part
    for rel_id, rel in main_part.rels.items():
        if rel.reltype == RT.FOOTNOTES:
            fn_part = rel.target_part
            return etree.fromstring(fn_part.blob)

    # Создаём часть с нуля
    partname = PackURI("/word/footnotes.xml")
    content_type = CT.WML_FOOTNOTES
    fn_element = etree.fromstring(FOOTNOTES_XML_TEMPLATE.encode("utf-8"))
    blob = etree.tostring(
        fn_element, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    fn_part = Part(partname, content_type, blob, main_part.package)
    main_part.relate_to(fn_part, RT.FOOTNOTES)
    return fn_element


def _flush_footnotes(doc):
    """Записывает накопленный элемент ``<w:footnotes>`` обратно в часть
    word/footnotes.xml после того, как все ``<w:footnote>`` добавлены.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from lxml import etree

    main_part = doc.part
    for rel_id, rel in main_part.rels.items():
        if rel.reltype == RT.FOOTNOTES:
            fn_part = rel.target_part
            fn_part._blob = etree.tostring(  # type: ignore[attr-defined]
                _FOOTNOTES_ROOT, xml_declaration=True, encoding="UTF-8",
                standalone=True,
            )
            return


_FOOTNOTES_ROOT = None  # <w:footnotes> элемент; заполняется в setup_footnotes_part


def setup_footnotes_part(doc):
    """Подготавливает хранилище сносок. Вызывать один раз на документ перед
    началом рендеринга body, чтобы идентификаторы сносок (id=1, 2, ...)
    соответствовали порядку появления маркеров в тексте.
    """
    global _FOOTNOTES_ROOT
    _ensure_footnote_styles(doc)
    _FOOTNOTES_ROOT = _ensure_footnotes_part(doc)


def add_footnote_reference(paragraph, source_num, *, base_size=BODY_SIZE):
    """Вставляет в параграф маркер сноски (верхний индекс) и регистрирует
    соответствующий ``<w:footnote>`` в footnotes-part. Содержимое сноски —
    библиографическая запись источника №``source_num`` из ``SOURCES``.
    """
    global _FOOTNOTES_ROOT
    if _FOOTNOTES_ROOT is None:
        setup_footnotes_part(paragraph.part.document if hasattr(paragraph, "part") else None)

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    # Нумерация сносок: идентификатор footnote = порядковому номеру её
    # появления в документе (1, 2, 3, ...). Привязка к номеру источника
    # сохраняется через содержимое тела сноски.
    fn_id = _next_footnote_id()
    USED_FOOTNOTE_IDS.add(source_num)

    # 1. Маркер в теле: <w:r><w:rPr><w:rStyle val="FootnoteReference"/></w:rPr>
    #                   <w:footnoteReference w:id="N"/></w:r>
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "FootnoteReference")
    rPr.append(rStyle)
    r.append(rPr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fn_id))
    r.append(ref)
    paragraph._p.append(r)

    # 2. Тело сноски — полная запись по ГОСТ. Если источник не найден,
    # оставляем заглушку, чтобы разработчик увидел незакрытую ссылку.
    citation = SOURCES.get(source_num)
    if citation is None:
        citation = f"[Источник {source_num} не найден в списке использованных источников]"

    # Сноска по ГОСТ начинается с номера (он же «знак сноски»), затем точка,
    # затем текст библ. записи. Word сам проставит верхний индекс за счёт стиля
    # FootnoteReference у первого run'а.
    fn = OxmlElement("w:footnote")
    fn.set(qn("w:id"), str(fn_id))

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "FootnoteText")
    pPr.append(pStyle)
    p.append(pPr)

    # run 1 — номер сноски верхним индексом
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    rStyle1 = OxmlElement("w:rStyle")
    rStyle1.set(qn("w:val"), "FootnoteReference")
    rPr1.append(rStyle1)
    r1.append(rPr1)
    ref1 = OxmlElement("w:footnoteRef")
    r1.append(ref1)
    p.append(r1)

    # run 2 — пробел + библиографическая запись.
    # Шрифт и кегль задаём явно (TNR 14), чтобы не зависело от того,
    # подхватит ли Word rPr из стиля FootnoteText.
    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    rFonts2 = OxmlElement("w:rFonts")
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts2.set(qn(k), BODY_FONT)
    rPr2.append(rFonts2)
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "28")
    rPr2.append(sz2)
    szCs2 = OxmlElement("w:szCs")
    szCs2.set(qn("w:val"), "28")
    rPr2.append(szCs2)
    r2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.set(qn("xml:space"), "preserve")
    t2.text = " " + citation
    r2.append(t2)
    p.append(r2)

    fn.append(p)
    _FOOTNOTES_ROOT.append(fn)


_FOOTNOTE_COUNTER = 0


def _next_footnote_id():
    global _FOOTNOTE_COUNTER
    _FOOTNOTE_COUNTER += 1
    return _FOOTNOTE_COUNTER


def parse_sources_from_md(md_text):
    """Парсит раздел «Список использованных источников» и заполняет SOURCES.

    Формат записи в md:
        ``N. Авторы. Заглавие. Год. URL: ... (дата обращения DD.MM.YYYY).``
    """
    lines = md_text.split("\n")
    in_sources = False
    current_num = None
    current_text_parts: "list[str]" = []

    def flush():
        if current_num is not None:
            SOURCES[current_num] = " ".join(current_text_parts).strip()

    for raw in lines:
        line = raw.rstrip()
        if re.match(r"^#\s+Список использованных источников\s*$", line):
            in_sources = True
            continue
        if not in_sources:
            continue
        # Начало новой записи «N. ...»
        m = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if m:
            flush()
            current_num = int(m.group(1))
            current_text_parts = [m.group(2).strip()]
            continue
        # Продолжение текущей записи (редко, обычно одна строка)
        if current_num is not None and line.strip():
            current_text_parts.append(line.strip())
    flush()


def add_block_formula(doc, latex):
    """Блочная формула: нативное OMML-уравнение Word, по центру."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.15
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    omml = _latex_to_omml(latex, display=True)
    if omml is None:
        r = p.add_run(latex_to_unicode(latex))
        run_set_font(r, MATH_FONT, size=BODY_SIZE, italic=True)
        return
    el = _omml_element(omml)
    if el is None:
        r = p.add_run(latex_to_unicode(latex))
        run_set_font(r, MATH_FONT, size=BODY_SIZE, italic=True)
        return
    p._p.append(el)


def add_inline_formula(paragraph, latex, *, base_size=BODY_SIZE):
    """Встроенная формула: нативное OMML-уравнение Word в текущем параграфе."""
    omml = _latex_to_omml(latex, display=False)
    if omml is None:
        r = paragraph.add_run(latex_to_unicode(latex))
        run_set_font(r, MATH_FONT, size=base_size, italic=True)
        return
    el = _omml_element(omml)
    if el is None:
        r = paragraph.add_run(latex_to_unicode(latex))
        run_set_font(r, MATH_FONT, size=base_size, italic=True)
        return
    paragraph._p.append(el)


def _compressed_image_stream(full_path, *, max_side=1800, jpeg_quality=82):
    """Масштабирует и пережимает изображение в JPEG-поток для встраивания в docx.
    Если изображение уже маленькое (<= max_side по длинной стороне) и в JPEG,
    возвращает исходный путь без изменений.
    """
    from PIL import Image
    import io

    try:
        with Image.open(full_path) as im:
            im.load()
            w, h = im.size
            longest = max(w, h)
            # решаем, нужна ли пересборка
            fmt = (im.format or "").upper()
            already_small = longest <= max_side
            already_jpeg = fmt in ("JPEG", "JPG")
            if already_small and already_jpeg:
                return full_path

            # масштабирование
            if longest > max_side:
                scale = max_side / longest
                new_size = (max(1, int(w * scale)), max(1, int(h * scale)))
                im = im.resize(new_size, Image.LANCZOS)

            # приводим к RGB с белой подложкой (для прозрачности)
            if im.mode in ("RGBA", "LA"):
                bg = Image.new("RGB", im.size, (255, 255, 255))
                alpha = im.split()[-1]
                bg.paste(im.convert("RGB"), mask=alpha)
                im = bg
            elif im.mode != "RGB":
                im = im.convert("RGB")

            buf = io.BytesIO()
            im.save(buf, format="JPEG", quality=jpeg_quality, optimize=True,
                    progressive=True)
            buf.seek(0)
            return buf
    except Exception:
        # при любой ошибке возвращаем оригинал
        return full_path


def add_image(doc, rel_path, *, caption_preceding=None):
    """Вставляет изображение, ширина ~ 14 см. Крупные PNG пережимаются в JPEG."""
    full_path = os.path.join(FIG_DIR, rel_path) if not os.path.isabs(rel_path) else rel_path
    if not os.path.exists(full_path):
        # Регистрируем плейсхолдер
        placeholders.append((caption_preceding, rel_path))
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        r = p.add_run(f"[ВСТАВИТЬ ИЗОБРАЖЕНИЕ: {caption_preceding or rel_path}]")
        run_set_font(r, BODY_FONT, size=BODY_SIZE, italic=True,
                     color=RGBColor(0xCC, 0x00, 0x00))
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run()
    try:
        # Узнаём реальные пропорции, чтобы ограничить высоту страницы
        target_w_cm = 14.0
        max_h_cm = 19.0  # оставляем место на подпись в пределах одного листа A4
        try:
            from PIL import Image
            with Image.open(full_path) as probe:
                pw, ph = probe.size
            aspect = ph / pw if pw else 0
        except Exception:
            aspect = 0
        if aspect and target_w_cm * aspect > max_h_cm:
            # высокое изображение — уменьшаем ширину так, чтобы высота = max_h_cm
            target_w_cm = max_h_cm / aspect
        payload = _compressed_image_stream(full_path)
        run.add_picture(payload, width=Cm(target_w_cm))
    except Exception as e:
        r = p.add_run(f"[ОШИБКА ВСТАВКИ: {rel_path} — {e}]")
        run_set_font(r, BODY_FONT, size=BODY_SIZE, italic=True,
                     color=RGBColor(0xCC, 0x00, 0x00))


def add_table(doc, header, rows):
    """header: list of strings, rows: list of list of strings."""
    n_cols = max(len(header), *[len(r) for r in rows]) if rows else len(header)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    # границы
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)

    def fill(cell, text, *, header_row=False):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # очищаем дефолтный абзац
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.15
        # Заголовочная строка по центру, данные — по левому краю. Но
        # жирного начертания не применяем нигде (A-18).
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if header_row else WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        add_inline_runs(p, text, base_size=TABLE_SIZE, base_bold=False)

    # заголовок
    for i, h in enumerate(header):
        if i < n_cols:
            fill(tbl.rows[0].cells[i], h, header_row=True)
    # данные
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < n_cols:
                fill(tbl.rows[1 + ri].cells[ci], cell_text)


# -----------------------------------------------------------------------------
# Основной парсер
# -----------------------------------------------------------------------------

RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
RE_IMAGE = re.compile(r"^!\[.*?\]\((.*?)\)\s*$")
RE_OLIST = re.compile(r"^(\d+)\.\s+(.*)$")
RE_BLIST = re.compile(r"^[-*]\s+(.*)$")
RE_BLOCK_MATH = re.compile(r"^\$\$(.*)\$\$\s*$")
RE_BLOCK_MATH_OPEN = re.compile(r"^\$\$(.*)$")
RE_BLOCK_MATH_CLOSE = re.compile(r"^(.*)\$\$\s*$")
RE_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
RE_TABLE_SEP = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")
RE_HR = re.compile(r"^---+\s*$")
RE_CODE_FENCE = re.compile(r"^```(\w*)\s*$")


def split_table_row(line):
    # "| a | b | c |"
    inner = line.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


def convert():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()
    lines = md_text.split("\n")

    # Разбираем список использованных источников — из него берутся тексты
    # подстраничных сносок для маркеров **[N]**.
    parse_sources_from_md(md_text)

    doc = setup_document()

    # Подготавливаем footnotes-part до начала рендеринга, чтобы каждый
    # обнаруженный маркер **[N]** уже мог добавить сноску.
    setup_footnotes_part(doc)

    # Авто-TOC не вставляем: по запросу автора оглавление собирается вручную
    # в Word через «Ссылки → Оглавление». В заголовках остаётся w:outlineLvl
    # (см. add_heading), в settings.xml — w:updateFields=true; этого достаточно,
    # чтобы ручная сборка оглавления в Word работала корректно.

    # Две зрительно пустые страницы в начале документа — место для титульного
    # листа и задания (оба заполняются автором в Word после сборки). Номера на
    # этих страницах не отображаются (раздел 1 не содержит футера).
    add_filler_page(doc)
    add_filler_page(doc)

    # Состояние для вставки разрыва раздела и пустой страницы под оглавление
    # после реферата. Логика: как только после реферата встречается следующий
    # заголовок (обычно «Термины и определения»), перед ним вставляется
    # разрыв раздела — начинается раздел 2 с нумерацией страниц с 4, и
    # добавляется пустая страница под автооглавление.
    referat_seen = False
    section2_started = False

    i = 0
    last_figure_caption = None  # последняя "подпись" перед изображением

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Пустая строка
        if stripped == "":
            i += 1
            continue

        # Горизонтальная черта -> пропуск (разделитель глав)
        if RE_HR.match(stripped):
            i += 1
            continue

        # Код-блок
        m = RE_CODE_FENCE.match(stripped)
        if m:
            lang = m.group(1)
            code_lines = []
            i += 1
            while i < len(lines) and not RE_CODE_FENCE.match(lines[i]):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # пропустим закрывающий ```
            add_code_block(doc, code_lines, lang=lang)
            continue

        # Блочная формула $$...$$
        m = RE_BLOCK_MATH.match(stripped)
        if m:
            add_block_formula(doc, m.group(1))
            i += 1
            continue

        # Многострочная блочная формула $$\n ... \n$$
        if stripped.startswith("$$") and not stripped.endswith("$$"):
            math_content = [stripped[2:]]
            i += 1
            while i < len(lines) and "$$" not in lines[i]:
                math_content.append(lines[i])
                i += 1
            if i < len(lines):
                last = lines[i]
                idx = last.find("$$")
                math_content.append(last[:idx])
                i += 1
            add_block_formula(doc, " ".join(math_content))
            continue

        # Заголовок
        m = RE_HEADING.match(stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            is_chapter = (level == 1)

            # После реферата — разрыв раздела, переход на пустую страницу под
            # автооглавление, старт нумерации страниц с 4.
            if referat_seen and not section2_started:
                start_numbered_section(doc, start_page=4, page_number_size=Pt(14))
                # Пустая страница под автооглавление: параграф сидит на
                # странице 4 в начале раздела 2, следующий заголовок
                # (обычно «Термины и определения») имеет pageBreakBefore и
                # переходит на страницу 5.
                p_toc = doc.add_paragraph()
                p_toc.paragraph_format.first_line_indent = Cm(0)
                section2_started = True

            add_heading(doc, level, text, is_chapter=is_chapter)

            if text.strip().lower() == "реферат":
                referat_seen = True

            i += 1
            continue

        # Изображение
        m = RE_IMAGE.match(stripped)
        if m:
            path = m.group(1)
            add_image(doc, path, caption_preceding=last_figure_caption)
            # Подпись рисунка ставим ПОД картинкой, по центру
            if last_figure_caption is not None:
                add_figure_caption(doc, last_figure_caption)
            last_figure_caption = None
            i += 1
            continue

        # Таблица
        if RE_TABLE_ROW.match(line) and i + 1 < len(lines) and RE_TABLE_SEP.match(lines[i + 1]):
            header = split_table_row(line)
            i += 2  # пропустим разделитель
            rows = []
            while i < len(lines) and RE_TABLE_ROW.match(lines[i]):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # Нумерованный список
        m = RE_OLIST.match(stripped)
        if m:
            num = m.group(1)
            text = m.group(2)
            add_list_item(doc, text, number=num)
            i += 1
            continue

        # Маркированный список
        m = RE_BLIST.match(stripped)
        if m:
            text = m.group(1)
            add_list_item(doc, text, bullet=True)
            i += 1
            continue

        # Подпись рисунка: отложим — вставим ПОД картинкой после её рендера
        if stripped.startswith("Рисунок "):
            last_figure_caption = stripped
            i += 1
            continue

        # Подпись таблицы: рендерим СВЕРХУ и СЛЕВА
        if stripped.startswith("Таблица "):
            add_table_caption(doc, stripped)
            i += 1
            continue

        # Обычный абзац
        add_paragraph(doc, stripped)
        last_figure_caption = None
        i += 1

    # Записываем накопленные сноски обратно в word/footnotes.xml.
    _flush_footnotes(doc)

    # В смонтированную рабочую папку python-docx не всегда может записать
    # .docx напрямую: если файл уже открыт в Word на хосте, он заблокирован
    # для перезаписи и мы получим FileNotFoundError. В этом случае падаем
    # с понятным сообщением — никаких запасных имён типа diplom_new.docx
    # создавать не нужно: выходной файл всегда один, diplom.docx.
    # На смонтированной Cowork-папке прямое открытие diplom.docx на запись
    # иногда возвращает FileNotFoundError (особенность FUSE), хотя создание
    # файла с другим именем в той же директории проходит штатно. Надёжный
    # путь — сохранить в соседний файл diplom_pre.docx и атомарно переименовать.
    out_dir = os.path.dirname(OUT_PATH) or "."
    staging_path = os.path.join(out_dir, "diplom_pre.docx")
    try:
        doc.save(staging_path)
        try:
            if os.path.exists(OUT_PATH):
                os.remove(OUT_PATH)
        except OSError:
            pass
        try:
            os.rename(staging_path, OUT_PATH)
        except OSError as err:
            raise RuntimeError(
                f"Не удалось переименовать {staging_path} в {OUT_PATH}: "
                f"{err.__class__.__name__}. Возможно, diplom.docx открыт в Word — "
                f"закройте файл и запустите сборку снова."
            ) from err
    except Exception:
        # Если не вышло — подчистить staging-файл
        try:
            if os.path.exists(staging_path):
                os.remove(staging_path)
        except OSError:
            pass
        raise
    return OUT_PATH


def add_paragraph_centered(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.keep_with_next = True
    add_inline_runs(p, text)


def add_figure_caption(doc, text):
    """Подпись под рисунком: по центру, одна строка, без отступа первой строки."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.15
    pf.space_before = Pt(3)
    pf.space_after = Pt(12)
    add_inline_runs(p, text)


def add_table_caption(doc, text):
    """Подпись над таблицей: слева, без отступа первой строки, прижата к таблице."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.15
    pf.space_before = Pt(12)
    pf.space_after = Pt(3)
    pf.keep_with_next = True
    add_inline_runs(p, text)


if __name__ == "__main__":
    out = convert()
    print(f"Saved: {out}")
    print(f"\nPlaceholders ({len(placeholders)}):")
    for cap, path in placeholders:
        print(f"  - {cap} -> {path}")
    print(f"\nSources parsed: {len(SOURCES)}")
    print(f"Footnote markers resolved: {sorted(USED_FOOTNOTE_IDS)}")
    missing = sorted(USED_FOOTNOTE_IDS - set(SOURCES.keys()))
    if missing:
        print(f"WARNING: markers **[N]** without matching source entry: {missing}")
