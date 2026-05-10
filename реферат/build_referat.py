#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Конвертер referat.md -> referat.docx.

Оформление по методичке МГИМО:
    - шрифт Times New Roman 14pt, междустрочный интервал 1.5
    - поля: левое 25 мм, правое 10 мм, верхнее и нижнее 15 мм
    - абзацный отступ 1.25 см, выравнивание по ширине
    - заголовки глав (Введение, Глава N, Заключение, Список) —
      ПРОПИСНЫЕ, по центру, с новой страницы
    - заголовки параграфов (1.1, 1.2, ...) — с абзацного отступа,
      строчные (кроме первой буквы), без жирного
    - подстрочные сноски: 10 pt, интервал 1.0, сквозная нумерация
    - нумерация страниц: с третьей страницы (после двух пустых,
      где будут размещены титульный лист и содержание),
      правый нижний угол, 12 pt
"""

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt
from lxml import etree


ROOT = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(ROOT, "referat.md")
OUT_PATH = os.path.join(ROOT, "..", "referat.docx")


BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(14)
FOOTNOTE_SIZE = Pt(10)
PAGE_NUM_SIZE = Pt(12)


# Структурные разделы (без нумерации, прописными, по центру, с новой страницы)
STRUCTURAL_SECTIONS = {
    "введение",
    "заключение",
    "список использованных источников и литературы",
}


# Источники: парсятся из секции "Список использованных источников"
SOURCES: "dict[int, str]" = {}


# -----------------------------------------------------------------------------
# Стили текста, шрифт по умолчанию
# -----------------------------------------------------------------------------

def set_default_style(doc):
    """Базовый стиль Normal: TNR 14, выравнивание по ширине, интервал 1.5,
    абзацный отступ 1.25 см, отбивки 0.
    """
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:cs"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)

    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5


def set_page_setup(section):
    """Поля по методичке МГИМО: левое 25 мм, правое 10 мм,
    верхнее и нижнее 15 мм. Размер бумаги — A4.
    """
    section.left_margin = Mm(25)
    section.right_margin = Mm(10)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.page_height = Mm(297)
    section.page_width = Mm(210)


# -----------------------------------------------------------------------------
# Маркеры сносок
# -----------------------------------------------------------------------------

# Маркер сноски: [N] или [N, M, K] — по идентификаторам источников
# Допускаем как одиночные [N], так и составные через запятую: [3, 5]
FOOTNOTE_RE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")


def parse_footnote_token(token):
    """Из текста '3, 5, 17' возвращает список int [3, 5, 17]."""
    return [int(x.strip()) for x in token.split(",")]


def split_text_with_footnotes(text):
    """Разбивает текст на список фрагментов: ('text', s) или ('foot', [n1, n2, ...])."""
    out = []
    pos = 0
    for m in FOOTNOTE_RE.finditer(text):
        if m.start() > pos:
            out.append(("text", text[pos:m.start()]))
        out.append(("foot", parse_footnote_token(m.group(1))))
        pos = m.end()
    if pos < len(text):
        out.append(("text", text[pos:]))
    return out


# -----------------------------------------------------------------------------
# Сноски (footnote-part)
# -----------------------------------------------------------------------------

FOOTNOTES_XML_TEMPLATE = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
</w:footnotes>
"""


def ensure_footnote_styles(doc):
    """Регистрирует стили FootnoteReference (верхний индекс) и
    FootnoteText (10 pt, интервал 1.0).
    """
    styles_el = doc.styles.element

    def has_style(style_id):
        for s in styles_el.findall(qn("w:style")):
            if s.get(qn("w:styleId")) == style_id:
                return True
        return False

    if not has_style("FootnoteReference"):
        s = OxmlElement("w:style")
        s.set(qn("w:type"), "character")
        s.set(qn("w:styleId"), "FootnoteReference")
        name = OxmlElement("w:name")
        name.set(qn("w:val"), "footnote reference")
        s.append(name)
        rPr = OxmlElement("w:rPr")
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "superscript")
        rPr.append(vert)
        s.append(rPr)
        styles_el.append(s)

    if not has_style("FootnoteText"):
        s = OxmlElement("w:style")
        s.set(qn("w:type"), "paragraph")
        s.set(qn("w:styleId"), "FootnoteText")
        name = OxmlElement("w:name")
        name.set(qn("w:val"), "footnote text")
        s.append(name)
        pPr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), "0")
        pPr.append(ind)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)
        s.append(pPr)
        rPr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), BODY_FONT)
        rfonts.set(qn("w:hAnsi"), BODY_FONT)
        rfonts.set(qn("w:cs"), BODY_FONT)
        rPr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")  # 10 pt
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "20")
        rPr.append(szCs)
        s.append(rPr)
        styles_el.append(s)


def ensure_footnotes_part(doc):
    """Создаёт word/footnotes.xml, если его нет. Возвращает корневой элемент."""
    from docx.opc.constants import CONTENT_TYPE as CT
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    main_part = doc.part
    for rel_id, rel in main_part.rels.items():
        if rel.reltype == RT.FOOTNOTES:
            fn_part = rel.target_part
            return etree.fromstring(fn_part.blob)

    partname = PackURI("/word/footnotes.xml")
    content_type = CT.WML_FOOTNOTES
    fn_element = etree.fromstring(FOOTNOTES_XML_TEMPLATE.encode("utf-8"))
    blob = etree.tostring(
        fn_element, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    fn_part = Part(partname, content_type, blob, main_part.package)
    main_part.relate_to(fn_part, RT.FOOTNOTES)
    return fn_element


_FOOTNOTES_ROOT = None
_NEXT_FN_ID = 1


def setup_footnotes(doc):
    global _FOOTNOTES_ROOT, _NEXT_FN_ID
    ensure_footnote_styles(doc)
    _FOOTNOTES_ROOT = ensure_footnotes_part(doc)
    _NEXT_FN_ID = 1


def flush_footnotes(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    main_part = doc.part
    for rel_id, rel in main_part.rels.items():
        if rel.reltype == RT.FOOTNOTES:
            fn_part = rel.target_part
            fn_part._blob = etree.tostring(
                _FOOTNOTES_ROOT, xml_declaration=True, encoding="UTF-8",
                standalone=True,
            )
            return


def add_footnote(paragraph, source_num):
    """Вставляет в параграф маркер сноски и регистрирует её содержимое."""
    global _NEXT_FN_ID
    fn_id = _NEXT_FN_ID
    _NEXT_FN_ID += 1

    # 1. Маркер в теле абзаца
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "FootnoteReference")
    rPr.append(rStyle)
    r.append(rPr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fn_id))
    r.append(ref)
    paragraph._p.append(r)

    # 2. Содержимое сноски
    citation = SOURCES.get(source_num,
                           f"[Источник {source_num} не найден]")

    fn = OxmlElement("w:footnote")
    fn.set(qn("w:id"), str(fn_id))

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "FootnoteText")
    pPr.append(pStyle)
    p.append(pPr)

    # Маркер в начале сноски (верхний индекс)
    r0 = OxmlElement("w:r")
    rPr0 = OxmlElement("w:rPr")
    rStyle0 = OxmlElement("w:rStyle")
    rStyle0.set(qn("w:val"), "FootnoteReference")
    rPr0.append(rStyle0)
    r0.append(rPr0)
    fnref = OxmlElement("w:footnoteRef")
    r0.append(fnref)
    p.append(r0)

    # Пробел и текст библиографической записи
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rPr1.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr1.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "20")
    rPr1.append(szCs)
    r1.append(rPr1)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = " " + citation
    r1.append(t)
    p.append(r1)

    fn.append(p)
    _FOOTNOTES_ROOT.append(fn)


# -----------------------------------------------------------------------------
# Параграфы и заголовки
# -----------------------------------------------------------------------------

def add_run(paragraph, text, *, size=BODY_SIZE, bold=False, italic=False,
            font=BODY_FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)
    rfonts.set(qn("w:eastAsia"), font)
    return run


def add_paragraph_with_footnotes(doc, text, *, indent=Cm(1.25),
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Добавляет абзац основного текста с автоматической раскруткой [N]-маркеров
    в подстрочные сноски.
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = indent
    p.paragraph_format.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Если в строке есть жирный markdown **...**, обрабатываем
    parts = split_text_with_footnotes(text)
    for kind, content in parts:
        if kind == "text":
            # Простой текст — без обработки markdown bold (текст уже без него)
            add_run(p, content)
        elif kind == "foot":
            for n in content:
                add_footnote(p, n)
    return p


def add_structural_heading(doc, text, *, page_break=True):
    """ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ —
    прописные, по центру, с новой страницы, без жирного, outlineLvl=0.
    """
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, text.upper(), bold=True)
    # outlineLvl для возможного оглавления
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "0")
    pPr.append(outline)


def add_chapter_heading(doc, number, title, *, page_break=True):
    """Глава N. Название — прописными, по центру, с новой страницы."""
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    full = f"ГЛАВА {number}. {title.upper()}"
    add_run(p, full, bold=True)
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "0")
    pPr.append(outline)


def add_section_heading(doc, number, title):
    """Параграф (1.1, 1.2, 2.1, ...) — с абзацного отступа, строчные
    (кроме первой буквы), без жирного.
    """
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, f"{number}. {title}", bold=True)
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "1")
    pPr.append(outline)


def add_subsection_heading(doc, title):
    """Подзаголовок третьего уровня (1. Источники / 2. Литература) —
    с красной строки, выделенный жирным, в списке литературы.
    """
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, title, bold=True)
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "2")
    pPr.append(outline)


# -----------------------------------------------------------------------------
# Нумерация страниц (с третьей страницы, правый нижний угол, 12 pt)
# -----------------------------------------------------------------------------

def setup_page_numbers(doc):
    """Реализуем модель: первая секция — 2 пустые страницы без номеров;
    вторая секция — основной текст со сквозной нумерацией, начинающейся с 3.
    """
    pass  # вызывается извне после построения секций


# -----------------------------------------------------------------------------
# Парсер референса
# -----------------------------------------------------------------------------

def parse_sources(md_text):
    """Извлекает SOURCES = {N: 'библиографическая запись'} из секции
    «Список использованных источников и литературы».
    """
    # Ищем блок начиная с заголовка списка источников
    lines = md_text.split("\n")
    in_list = False
    out = {}
    for line in lines:
        s = line.strip()
        if not in_list:
            if s.startswith("# СПИСОК"):
                in_list = True
            continue
        # Внутри списка: ищем строки вида "1. ..." или "12. ..."
        m = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m:
            n = int(m.group(1))
            entry = m.group(2)
            out[n] = entry
    return out


# -----------------------------------------------------------------------------
# Главная функция
# -----------------------------------------------------------------------------

def build():
    with open(MD_PATH, encoding="utf-8") as f:
        md_text = f.read()

    global SOURCES
    SOURCES = parse_sources(md_text)
    print(f"Найдено источников: {len(SOURCES)}", file=sys.stderr)

    doc = Document()
    set_default_style(doc)

    # Первая секция: 2 пустые страницы (для титульника и оглавления)
    section_blank = doc.sections[0]
    set_page_setup(section_blank)
    # Без номеров страниц
    sect_pr = section_blank._sectPr
    titlePg = sect_pr.find(qn("w:titlePg"))

    # Добавляем 2 пустых параграфа с разрывами страниц
    p1 = doc.add_paragraph()
    p1.add_run().add_break(WD_BREAK.PAGE)
    p2 = doc.add_paragraph()
    p2.add_run().add_break(WD_BREAK.PAGE)

    # Запускаем новую секцию для основного содержимого
    new_sect_para = doc.add_paragraph()
    new_sect_pPr = new_sect_para._p.get_or_add_pPr()
    new_sectPr = OxmlElement("w:sectPr")
    new_sect_pPr.append(new_sectPr)

    # Скопируем настройки страницы в секцию пустых страниц
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "11906")  # A4 ширина в twips
    pgSz.set(qn("w:h"), "16838")
    new_sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), str(int(Mm(15).twips)))
    pgMar.set(qn("w:right"), str(int(Mm(10).twips)))
    pgMar.set(qn("w:bottom"), str(int(Mm(15).twips)))
    pgMar.set(qn("w:left"), str(int(Mm(25).twips)))
    pgMar.set(qn("w:header"), "720")
    pgMar.set(qn("w:footer"), "720")
    pgMar.set(qn("w:gutter"), "0")
    new_sectPr.append(pgMar)
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    new_sectPr.append(type_el)

    setup_footnotes(doc)

    # Парсим основной текст и строим документ
    parse_and_build(doc, md_text)

    # Применяем нумерацию страниц во второй (последней) секции
    apply_page_numbering(doc)

    flush_footnotes(doc)
    doc.save(OUT_PATH)
    print(f"Сохранено: {OUT_PATH}", file=sys.stderr)


def parse_and_build(doc, md_text):
    """Парсит md и добавляет в документ. Учитывает заголовки # / ## / ###
    и обычные абзацы.
    """
    lines = md_text.split("\n")
    i = 0
    in_sources = False
    first_h1 = True
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # Заголовок первого уровня
        if s.startswith("# "):
            title = s[2:].strip()
            t_lower = title.lower()
            page_break = not first_h1
            first_h1 = False
            if t_lower in STRUCTURAL_SECTIONS:
                add_structural_heading(doc, title, page_break=page_break)
                if t_lower == "список использованных источников и литературы":
                    in_sources = True
                else:
                    in_sources = False
            elif t_lower.startswith("глава"):
                m = re.match(r"глава\s+(\d+)\.\s+(.+)", t_lower)
                if m:
                    num = m.group(1)
                    # Сохраняем оригинальный регистр
                    orig = re.match(r"Глава\s+\d+\.\s+(.+)", title)
                    chap_title = orig.group(1) if orig else title
                    add_chapter_heading(doc, num, chap_title, page_break=page_break)
                else:
                    add_structural_heading(doc, title, page_break=page_break)
                in_sources = False
            else:
                add_structural_heading(doc, title, page_break=page_break)
                in_sources = False
            i += 1
            continue

        # Заголовок второго уровня — параграф NN.NN
        if s.startswith("## "):
            title = s[3:].strip()
            m = re.match(r"^(\d+\.\d+)\.\s+(.+)", title)
            if m:
                add_section_heading(doc, m.group(1), m.group(2))
            else:
                add_subsection_heading(doc, title)
            i += 1
            continue

        # Заголовок третьего уровня
        if s.startswith("### "):
            title = s[4:].strip()
            add_subsection_heading(doc, title)
            i += 1
            continue

        # В списке источников строки вида "1. ..." — нумерованный список
        if in_sources:
            m = re.match(r"^(\d+)\.\s+(.+)$", s)
            if m:
                # Источник — отдельный абзац с висячей нумерацией
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0.75)
                # hanging indent имитируем через табуляцию
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                add_run(p, f"{m.group(1)}. {m.group(2)}")
                i += 1
                continue

        # Список (нумерованный) — пункт "1. ..."
        m = re.match(r"^(\d+)\.\s+(.+)$", s)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # Включаем номер прямо в текст
            text = f"{m.group(1)}) {m.group(2)}"
            parts = split_text_with_footnotes(text)
            for kind, content in parts:
                if kind == "text":
                    add_run(p, content)
                else:
                    for n in content:
                        add_footnote(p, n)
            i += 1
            continue

        # Обычный абзац
        add_paragraph_with_footnotes(doc, s)
        i += 1


# -----------------------------------------------------------------------------
# Нумерация страниц
# -----------------------------------------------------------------------------

def apply_page_numbering(doc):
    """Двусекционная схема:
    - Первая секция (титульник + содержание) — 2 пустые страницы, без номеров.
    - Вторая секция (введение и далее) — нумерация начинается с 3,
      правый нижний угол, 12 pt.
    """
    sections = doc.sections
    # Включаем разные header/footer для первой страницы (не нужно)
    # Скрываем номер на первых двух страницах через настройки первой секции.

    # Первая секция: убираем footer (без номеров)
    first_sect = sections[0]
    # Гарантируем, что у первой секции нет номеров — оставляем footer пустым
    first_footer = first_sect.footer
    first_footer.is_linked_to_previous = False
    # Очищаем footer
    for p in first_footer.paragraphs:
        for r in list(p.runs):
            r._element.getparent().remove(r._element)

    # Вторая секция (последняя): нумерация с 3
    if len(sections) >= 2:
        last = sections[-1]
        last_footer = last.footer
        last_footer.is_linked_to_previous = False

        # Очищаем существующие параграфы footer
        for p in list(last_footer.paragraphs):
            for r in list(p.runs):
                r._element.getparent().remove(r._element)

        # Добавляем параграф с PAGE field, выравнивание по правому краю
        if last_footer.paragraphs:
            p = last_footer.paragraphs[0]
        else:
            p = last_footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)

        run = p.add_run()
        run.font.name = BODY_FONT
        run.font.size = PAGE_NUM_SIZE
        rpr = run._element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), BODY_FONT)
        rfonts.set(qn("w:hAnsi"), BODY_FONT)
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")  # 12 pt
        rpr.append(sz)

        # Создаём поле PAGE
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run._element.append(instr)

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._element.append(fld_sep)

        # Поле для отображения значения
        t = OxmlElement("w:t")
        t.text = "3"
        run._element.append(t)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._element.append(fld_end)

        # Устанавливаем pgNumType для второй секции с начальным значением 3
        sect_pr = last._sectPr
        # удаляем существующий pgNumType
        existing = sect_pr.find(qn("w:pgNumType"))
        if existing is not None:
            sect_pr.remove(existing)
        pgNumType = OxmlElement("w:pgNumType")
        pgNumType.set(qn("w:start"), "3")
        sect_pr.append(pgNumType)

        # Также для первой секции: указываем, что нумерация будет начата
        # позже. Но эстетически проще убрать отображение в footer первой
        # секции (что мы и сделали выше).


if __name__ == "__main__":
    build()
